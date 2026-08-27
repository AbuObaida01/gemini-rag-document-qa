from pathlib import Path
from typing import Any

from docx import Document

from app.services.extraction.base import BaseExtractor


class DOCXExtractor(BaseExtractor):
    """
    Extract text from Microsoft Word DOCX files.

    Extracts:
    - Paragraphs
    - Tables
    """

    def extract(self, file_path: Path) -> dict[str, Any]:
        if not file_path.exists():
            raise FileNotFoundError(
                f"File not found: {file_path}"
            )

        if file_path.suffix.lower() != ".docx":
            raise ValueError(
                "DOCXExtractor can only process DOCX files."
            )

        document = Document(file_path)

        sections = []

        # Extract paragraphs
        for paragraph_number, paragraph in enumerate(
            document.paragraphs,
            start=1,
        ):
            text = paragraph.text.strip()

            if not text:
                continue

            sections.append(
                {
                    "text": text,
                    "metadata": {
                        "type": "paragraph",
                        "paragraph": paragraph_number,
                    },
                }
            )

        # Extract tables
        for table_number, table in enumerate(
            document.tables,
            start=1,
        ):
            rows = []

            for row in table.rows:
                cells = [
                    cell.text.strip()
                    for cell in row.cells
                ]

                rows.append(" | ".join(cells))

            table_text = "\n".join(rows).strip()

            if not table_text:
                continue

            sections.append(
                {
                    "text": table_text,
                    "metadata": {
                        "type": "table",
                        "table": table_number,
                    },
                }
            )

        if not sections:
            raise ValueError(
                "No extractable text was found in the DOCX file."
            )

        full_text = "\n\n".join(
            section["text"]
            for section in sections
        )

        return {
            "text": full_text,
            "sections": sections,
            "metadata": {
                "filename": file_path.name,
                "file_type": "docx",
                "paragraph_count": len(document.paragraphs),
                "table_count": len(document.tables),
            },
        }